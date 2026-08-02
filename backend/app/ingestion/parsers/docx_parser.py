from docx import Document
from docx.oxml.ns import qn
from .base_parser import BaseParser, ParsedPage

class DocxParser(BaseParser):
    """
      1. Page Break (ctrl+Enter hoặc Section Break trong Word) → chia trang cứng
      2. Nếu không có Page Break nào → toàn bộ file là 1 trang
    """

    # XML tag của Page Break trong .docx (thuộc nhánh w:br w:type="page")
    _PAGE_BREAK_QNAME = qn("w:br")

    def _paragraph_has_page_break(self, para) -> bool:
        """Trả về True nếu paragraph chứa dấu ngắt trang (Page Break)"""
        # Trong XML, page break nằm trong <w:br w:type="page">
        for br in para._element.findall(f".//{self._PAGE_BREAK_QNAME}"):
            br_type = br.get(qn("w:type"))
            if br_type == "page":
                return True
        return False

    def _extract_table_text(self, table) -> str:
        """Chuyển bảng Word sang dạng text pipe-separated đơn giản"""
        lines = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
        return "\n".join(lines)

    def parse(self, file_path: str) -> list[ParsedPage]:
        doc = Document(file_path)

        # Lấy toàn bộ nội dung body theo thứ tự (paragraphs + tables xen kẽ)
        # doc.element.body chứa tất cả các phần tử XML: w:p (para) và w:tbl (table)
        body = doc.element.body

        pages: list[ParsedPage] = []
        current_lines: list[str] = []
        page_num = 1

        # Tạo map từ XML element → python-docx object để tiện dùng
        para_map = {p._element: p for p in doc.paragraphs}
        table_map = {t._element: t for t in doc.tables}

        for child in body:
            tag = child.tag

            if tag == qn("w:p"):                    # Paragraph
                para = para_map.get(child)
                if para is None:
                    continue

                # Kiểm tra xem paragraph này có chứa Page Break không
                if self._paragraph_has_page_break(para):
                    # Lưu trang hiện tại trước khi sang trang mới
                    text = self.clean_text("\n".join(current_lines))
                    if len(text) >= 50:
                        pages.append(ParsedPage(
                            page_num=page_num,
                            text=text,
                            metadata={"char_count": len(text), "source": "docx"},
                        ))
                    page_num += 1
                    current_lines = []
                    # Vẫn lấy text phần còn lại của paragraph sau page break (nếu có)

                para_text = para.text.strip()
                if para_text:
                    current_lines.append(para_text)

            elif tag == qn("w:tbl"):                # Table
                table = table_map.get(child)
                if table is not None:
                    table_text = self._extract_table_text(table)
                    if table_text.strip():
                        current_lines.append(table_text)

        # Flush trang cuối cùng
        if current_lines:
            text = self.clean_text("\n".join(current_lines))
            if len(text) >= 50:
                pages.append(ParsedPage(
                    page_num=page_num,
                    text=text,
                    metadata={"char_count": len(text), "source": "docx"},
                ))

        # Nếu không có page break nào → 1 trang duy nhất đã được flush ở trên
        return pages
